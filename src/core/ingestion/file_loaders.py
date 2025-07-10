"""
File loader infrastructure for document ingestion.

Modular file loading system supporting multiple document formats
with extensible architecture for adding new file types.
"""

import abc
import asyncio
import base64
import csv
import io
import json
import re
from typing import Any, Dict, List, Optional, Tuple, Union

from ..utils.logger import get_logger

logger = get_logger(__name__)

# Global registry for file loaders
_FILE_LOADERS: Dict[str, type] = {}


class ParsedContent:
    """Container for parsed document content and metadata."""
    
    def __init__(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
        chunks: Optional[List[Dict[str, Any]]] = None,
        structure: Optional[Dict[str, Any]] = None,
        warnings: Optional[List[str]] = None,
    ):
        self.text = text
        self.metadata = metadata or {}
        self.chunks = chunks or []
        self.structure = structure or {}
        self.warnings = warnings or []


class BaseFileLoader(abc.ABC):
    """Abstract base class for file loaders."""
    
    @property
    @abc.abstractmethod
    def supported_types(self) -> List[str]:
        """Return list of supported file types."""
        pass
    
    @property
    @abc.abstractmethod
    def default_chunking_strategy(self) -> str:
        """Return the default chunking strategy for this file type."""
        pass
    
    @abc.abstractmethod
    async def load(self, content_bytes: bytes, file_name: str, **kwargs) -> ParsedContent:
        """Load and parse the file content."""
        pass
    
    @abc.abstractmethod
    async def extract_metadata(self, content_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """Extract metadata from the file."""
        pass
    
    def validate_content(self, content_bytes: bytes) -> bool:
        """Validate that the content can be processed by this loader."""
        return len(content_bytes) > 0


class PDFLoader(BaseFileLoader):
    """PDF file loader using PyMuPDF."""
    
    @property
    def supported_types(self) -> List[str]:
        return ["pdf"]
    
    @property
    def default_chunking_strategy(self) -> str:
        return "semantic"
    
    async def load(self, content_bytes: bytes, file_name: str, **kwargs) -> ParsedContent:
        """Load PDF content using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF is required for PDF processing. Install with: pip install PyMuPDF")
        
        try:
            # Open PDF from bytes
            doc = fitz.open(stream=content_bytes, filetype="pdf")
            
            text_content = []
            pages_metadata = []
            warnings = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text
                page_text = page.get_text()
                if not page_text.strip():
                    warnings.append(f"Page {page_num + 1} contains no extractable text")
                    continue
                
                text_content.append(page_text)
                
                # Extract page metadata
                page_metadata = {
                    "page_number": page_num + 1,
                    "text_length": len(page_text),
                    "bbox": page.rect,
                }
                pages_metadata.append(page_metadata)
            
            doc.close()
            
            # Combine all text
            full_text = "\n\n".join(text_content)
            
            # Create initial chunks by page
            chunks = []
            for i, page_text in enumerate(text_content):
                if page_text.strip():
                    chunks.append({
                        "text": page_text.strip(),
                        "chunk_type": "page",
                        "source_page": i + 1,
                        "metadata": pages_metadata[i]
                    })
            
            metadata = await self.extract_metadata(content_bytes, file_name)
            metadata.update({
                "page_count": len(doc),
                "text_length": len(full_text),
                "pages_with_text": len([p for p in text_content if p.strip()]),
            })
            
            return ParsedContent(
                text=full_text,
                metadata=metadata,
                chunks=chunks,
                structure={"pages": pages_metadata},
                warnings=warnings
            )
            
        except Exception as e:
            logger.error(f"Failed to load PDF: {str(e)}")
            raise
    
    async def extract_metadata(self, content_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """Extract PDF metadata."""
        try:
            import fitz
            
            doc = fitz.open(stream=content_bytes, filetype="pdf")
            metadata = doc.metadata
            doc.close()
            
            return {
                "title": metadata.get("title", ""),
                "author": metadata.get("author", ""),
                "subject": metadata.get("subject", ""),
                "creator": metadata.get("creator", ""),
                "producer": metadata.get("producer", ""),
                "creation_date": metadata.get("creationDate", ""),
                "modification_date": metadata.get("modDate", ""),
                "file_size": len(content_bytes),
                "format": "PDF",
            }
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata: {str(e)}")
            return {"file_size": len(content_bytes), "format": "PDF"}


class DOCXLoader(BaseFileLoader):
    """DOCX file loader using python-docx."""
    
    @property
    def supported_types(self) -> List[str]:
        return ["docx"]
    
    @property
    def default_chunking_strategy(self) -> str:
        return "paragraph"
    
    async def load(self, content_bytes: bytes, file_name: str, **kwargs) -> ParsedContent:
        """Load DOCX content using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            # Open document from bytes
            doc = Document(io.BytesIO(content_bytes))
            
            paragraphs = []
            chunks = []
            warnings = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue
                
                paragraphs.append(text)
                
                # Create chunk for each paragraph
                chunks.append({
                    "text": text,
                    "chunk_type": "paragraph",
                    "paragraph_index": i,
                    "style": paragraph.style.name if paragraph.style else "Normal",
                })
            
            # Extract tables
            table_texts = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                # Convert table to text
                table_text = "\n".join(["\t".join(row) for row in table_data])
                table_texts.append(table_text)
                
                chunks.append({
                    "text": table_text,
                    "chunk_type": "table",
                    "table_index": table_idx,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                })
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n" + "\n\n".join(table_texts)
            
            metadata = await self.extract_metadata(content_bytes, file_name)
            metadata.update({
                "paragraph_count": len(paragraphs),
                "table_count": len(table_texts),
                "text_length": len(full_text),
            })
            
            return ParsedContent(
                text=full_text,
                metadata=metadata,
                chunks=chunks,
                structure={
                    "paragraphs": len(paragraphs),
                    "tables": len(table_texts),
                },
                warnings=warnings
            )
            
        except Exception as e:
            logger.error(f"Failed to load DOCX: {str(e)}")
            raise
    
    async def extract_metadata(self, content_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """Extract DOCX metadata."""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content_bytes))
            props = doc.core_properties
            
            return {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "category": props.category or "",
                "comments": props.comments or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
                "file_size": len(content_bytes),
                "format": "DOCX",
            }
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {str(e)}")
            return {"file_size": len(content_bytes), "format": "DOCX"}


class TextLoader(BaseFileLoader):
    """Text file loader for TXT and MD files."""
    
    @property
    def supported_types(self) -> List[str]:
        return ["txt", "md"]
    
    @property
    def default_chunking_strategy(self) -> str:
        return "paragraph"
    
    async def load(self, content_bytes: bytes, file_name: str, **kwargs) -> ParsedContent:
        """Load text content with encoding detection."""
        try:
            import chardet
        except ImportError:
            logger.warning("chardet not available, falling back to utf-8")
            chardet = None
        
        # Detect encoding
        if chardet:
            encoding_result = chardet.detect(content_bytes)
            encoding = encoding_result.get("encoding", "utf-8")
            confidence = encoding_result.get("confidence", 0.0)
            
            if confidence < 0.7:
                logger.warning(f"Low confidence ({confidence:.2f}) in encoding detection: {encoding}")
        else:
            encoding = "utf-8"
        
        try:
            text = content_bytes.decode(encoding)
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            text = content_bytes.decode("utf-8", errors="replace")
            logger.warning("Used utf-8 with error replacement due to encoding issues")
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        
        # Create chunks
        chunks = []
        for i, paragraph in enumerate(paragraphs):
            chunks.append({
                "text": paragraph,
                "chunk_type": "paragraph",
                "paragraph_index": i,
            })
        
        # For Markdown, also detect headers
        if file_name.lower().endswith(('.md', '.markdown')):
            headers = re.findall(r'^(#{1,6})\s+(.+)$', text, re.MULTILINE)
            header_structure = [{"level": len(h[0]), "text": h[1]} for h in headers]
        else:
            header_structure = []
        
        metadata = await self.extract_metadata(content_bytes, file_name)
        metadata.update({
            "paragraph_count": len(paragraphs),
            "line_count": len(text.split("\n")),
            "character_count": len(text),
            "word_count": len(text.split()),
            "encoding": encoding,
        })
        
        return ParsedContent(
            text=text,
            metadata=metadata,
            chunks=chunks,
            structure={"headers": header_structure} if header_structure else {},
        )
    
    async def extract_metadata(self, content_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """Extract text file metadata."""
        return {
            "file_size": len(content_bytes),
            "format": "Markdown" if file_name.lower().endswith(('.md', '.markdown')) else "Text",
        }


class HTMLLoader(BaseFileLoader):
    """HTML file loader with text extraction."""
    
    @property
    def supported_types(self) -> List[str]:
        return ["html"]
    
    @property
    def default_chunking_strategy(self) -> str:
        return "semantic"
    
    async def load(self, content_bytes: bytes, file_name: str, **kwargs) -> ParsedContent:
        """Load HTML content and convert to text."""
        try:
            import html2text
        except ImportError:
            raise ImportError("html2text is required for HTML processing. Install with: pip install html2text")
        
        # Decode HTML content
        try:
            html_content = content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            html_content = content_bytes.decode("utf-8", errors="replace")
        
        # Convert HTML to text
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = True
        text_content = h.handle(html_content)
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in text_content.split("\n\n") if p.strip()]
        
        # Create chunks
        chunks = []
        for i, paragraph in enumerate(paragraphs):
            chunks.append({
                "text": paragraph,
                "chunk_type": "paragraph",
                "paragraph_index": i,
            })
        
        metadata = await self.extract_metadata(content_bytes, file_name)
        metadata.update({
            "paragraph_count": len(paragraphs),
            "original_html_size": len(html_content),
            "converted_text_size": len(text_content),
        })
        
        return ParsedContent(
            text=text_content,
            metadata=metadata,
            chunks=chunks,
        )
    
    async def extract_metadata(self, content_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """Extract HTML metadata."""
        try:
            html_content = content_bytes.decode("utf-8", errors="replace")
            
            # Extract title
            title_match = re.search(r'<title[^>]*>([^<]+)</title>', html_content, re.IGNORECASE)
            title = title_match.group(1).strip() if title_match else ""
            
            # Extract meta description
            desc_match = re.search(r'<meta[^>]*name=["\']description["\'][^>]*content=["\']([^"\']+)["\']', html_content, re.IGNORECASE)
            description = desc_match.group(1).strip() if desc_match else ""
            
            return {
                "title": title,
                "description": description,
                "file_size": len(content_bytes),
                "format": "HTML",
            }
        except Exception as e:
            logger.warning(f"Failed to extract HTML metadata: {str(e)}")
            return {"file_size": len(content_bytes), "format": "HTML"}


class JSONLoader(BaseFileLoader):
    """JSON file loader for structured data."""
    
    @property
    def supported_types(self) -> List[str]:
        return ["json"]
    
    @property
    def default_chunking_strategy(self) -> str:
        return "object"
    
    async def load(self, content_bytes: bytes, file_name: str, **kwargs) -> ParsedContent:
        """Load JSON content."""
        try:
            json_content = content_bytes.decode("utf-8")
            data = json.loads(json_content)
        except (UnicodeDecodeError, json.JSONDecodeError) as e:
            raise ValueError(f"Invalid JSON content: {str(e)}")
        
        # Convert JSON to readable text
        text_content = json.dumps(data, indent=2, ensure_ascii=False)
        
        # Create chunks based on structure
        chunks = []
        if isinstance(data, list):
            # Array of objects
            for i, item in enumerate(data):
                item_text = json.dumps(item, indent=2, ensure_ascii=False)
                chunks.append({
                    "text": item_text,
                    "chunk_type": "array_item",
                    "item_index": i,
                })
        elif isinstance(data, dict):
            # Object with properties
            for key, value in data.items():
                value_text = json.dumps({key: value}, indent=2, ensure_ascii=False)
                chunks.append({
                    "text": value_text,
                    "chunk_type": "object_property",
                    "property_key": key,
                })
        else:
            # Simple value
            chunks.append({
                "text": text_content,
                "chunk_type": "value",
            })
        
        metadata = await self.extract_metadata(content_bytes, file_name)
        metadata.update({
            "json_type": type(data).__name__,
            "item_count": len(data) if isinstance(data, (list, dict)) else 1,
            "text_length": len(text_content),
        })
        
        return ParsedContent(
            text=text_content,
            metadata=metadata,
            chunks=chunks,
            structure={"type": type(data).__name__, "keys": list(data.keys()) if isinstance(data, dict) else None},
        )
    
    async def extract_metadata(self, content_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """Extract JSON metadata."""
        return {
            "file_size": len(content_bytes),
            "format": "JSON",
        }


class CSVLoader(BaseFileLoader):
    """CSV file loader for tabular data."""
    
    @property
    def supported_types(self) -> List[str]:
        return ["csv"]
    
    @property
    def default_chunking_strategy(self) -> str:
        return "row"
    
    async def load(self, content_bytes: bytes, file_name: str, **kwargs) -> ParsedContent:
        """Load CSV content."""
        try:
            csv_content = content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            csv_content = content_bytes.decode("utf-8", errors="replace")
        
        # Parse CSV
        csv_reader = csv.reader(io.StringIO(csv_content))
        rows = list(csv_reader)
        
        if not rows:
            raise ValueError("Empty CSV file")
        
        # Assume first row is header
        headers = rows[0]
        data_rows = rows[1:]
        
        # Create text representation
        text_lines = []
        text_lines.append("Headers: " + ", ".join(headers))
        text_lines.append("")
        
        for i, row in enumerate(data_rows):
            row_text = " | ".join(f"{headers[j]}: {row[j]}" for j in range(min(len(headers), len(row))))
            text_lines.append(f"Row {i + 1}: {row_text}")
        
        text_content = "\n".join(text_lines)
        
        # Create chunks per row
        chunks = []
        for i, row in enumerate(data_rows):
            row_dict = {}
            for j, value in enumerate(row):
                if j < len(headers):
                    row_dict[headers[j]] = value
            
            row_text = json.dumps(row_dict, ensure_ascii=False)
            chunks.append({
                "text": row_text,
                "chunk_type": "csv_row",
                "row_index": i + 1,
                "row_data": row_dict,
            })
        
        metadata = await self.extract_metadata(content_bytes, file_name)
        metadata.update({
            "row_count": len(data_rows),
            "column_count": len(headers),
            "headers": headers,
            "text_length": len(text_content),
        })
        
        return ParsedContent(
            text=text_content,
            metadata=metadata,
            chunks=chunks,
            structure={"headers": headers, "row_count": len(data_rows)},
        )
    
    async def extract_metadata(self, content_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """Extract CSV metadata."""
        return {
            "file_size": len(content_bytes),
            "format": "CSV",
        }


# Register all file loaders
def register_file_loader(loader_class: type) -> None:
    """Register a file loader for specific file types."""
    loader = loader_class()
    for file_type in loader.supported_types:
        _FILE_LOADERS[file_type] = loader_class
        logger.debug(f"Registered {loader_class.__name__} for {file_type} files")


def get_file_loader(file_type: str) -> BaseFileLoader:
    """Get the appropriate file loader for a file type."""
    loader_class = _FILE_LOADERS.get(file_type.lower())
    if not loader_class:
        raise ValueError(f"No loader available for file type: {file_type}")
    
    return loader_class()


# Register all default loaders
register_file_loader(PDFLoader)
register_file_loader(DOCXLoader)
register_file_loader(TextLoader)
register_file_loader(HTMLLoader)
register_file_loader(JSONLoader)
register_file_loader(CSVLoader)

logger.info(f"Registered file loaders for: {list(_FILE_LOADERS.keys())}")